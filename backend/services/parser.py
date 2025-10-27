import os
import re
import json
from uuid import uuid4
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse
from docx import Document  # untuk baca DOCX
import pdfplumber  # untuk baca PDF
from pdf2image import convert_from_path
import aiofiles
from pathlib import Path
from fastapi import BackgroundTasks
import re
from PIL import Image
import pytesseract
import json
from pathlib import Path
from fastapi import FastAPI
from routes.keywords import router as keyword_router


UPLOAD_DIR = "uploads"
RESULTS_DIR = "results"
Path(RESULTS_DIR).mkdir(parents=True, exist_ok=True)
os.makedirs(UPLOAD_DIR, exist_ok=True)


def extract_text_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs])


def extract_text_from_pdf(file_path: str) -> str:
    text = ""

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

    # Fallback kalau text kosong → OCR (karena kemungkinan hasil scan)
    if not text.strip():
        print("[INFO] PDF kosong, fallback ke OCR mode...")
        images = convert_from_path(file_path, dpi=300)
        for img in images:
            text += pytesseract.image_to_string(img, lang="eng") + "\n"

    return text.strip()

SECTION_HEADERS = [

    # Work Experience Variants
    "work experience","Work Experiences", "Work Experience", "WORK EXPERIENCE", "experience", "Experience", "EXPERIENCE",
    "internship", "Internship", "INTERNSHIP", "internship experience", "Internship Experience", "INTERNSHIP EXPERIENCE",
    "professional experience", "Professional Experience", "PROFESSIONAL EXPERIENCE",
    "pengalaman", "Pengalaman", "PENGALAMAN",
    "pengalaman kerja", "Pengalaman Kerja", "PENGALAMAN KERJA",
    "pengalaman magang", "Pengalaman Magang", "PENGALAMAN MAGANG",
    "pengalaman profesional", "Pengalaman Profesional", "PENGALAMAN PROFESIONAL",
    "pengalaman industri", "Pengalaman Industri", "PENGALAMAN INDUSTRI",
    "Working Experience","WORK EXPERIENCE,","work experience"

    # Projects Variants
    "projects", "Projects", "PROJECTS", "project", "Project", "PROJECT",
    "personal projects", "Personal Projects", "PERSONAL PROJECTS",
    "portfolio", "Portfolio", "PORTFOLIO",
    "proyek", "Proyek", "PROYEK", "proyek pribadi", "Proyek Pribadi", "PROYEK PRIBADI",
    "PROJEK","PROJECT EXPERIENCE","Project Experience"

    # Education Variants
    "education", "Education","Education Level", "EDUCATION",
    "academic background", "Academic Background", "ACADEMIC BACKGROUND",
    "pendidikan", "Pendidikan", "PENDIDIKAN",
    "riwayat pendidikan", "Riwayat Pendidikan", "RIWAYAT PENDIDIKAN",
    "latihan akademik", "Latihan Akademik", "LATIHAN AKADEMIK",
    "EDUCATION AND QUALIFICATIONS", "Education and Qualifications"

    # Skills Variants
    "skills", "Skills", "SKILLS", "SKILSS", "Skilss", "SKILSS",
    "skill", "Skill", "SKILL",
    "technical skills", "Technical Skills", "TECHNICAL SKILLS",
    "soft skills", "Soft Skills", "SOFT SKILLS",
    "hard skills", "Hard Skills", "HARD SKILLS",
    "soft skill hard skill", "Soft Skill Hard Skill", "SOFT SKILL HARD SKILL",
    "keahlian", "Keahlian", "KEAHLIAN",
    "kemampuan", "Kemampuan", "KEMAMPUAN",
    "keterampilan", "Keterampilan", "KETERAMPILAN",

    # Certifications Variants
    "certifications", "Certifications", "CERTIFICATIONS",
    "certificates", "Certificates", "CERTIFICATES",
    "sertifikasi", "Sertifikasi", "SERTIFIKASI",
    "piagam", "Piagam", "PIAGAM",
    "lisensi", "Lisensi", "LISENSI",
    "pelatihan", "Pelatihan", "PELATIHAN",
    'pelatihan dan sertifikasi', 'Pelatihan dan Sertifikasi', 'PELATIHAN DAN SERTIFIKASI',
    "PELATIHAN/SERTIFIKASI","CERTIFICATION & TRAINING","CERTIFICATION","certification"

    # Achievements Variants
    "achievements", "Achievements", "ACHIEVEMENTS",
    "accomplishments", "Accomplishments", "ACCOMPLISHMENTS",
    "awards", "Awards", "AWARDS",
    "penghargaan", "Penghargaan", "PENGHARGAAN",
    "prestasi", "Prestasi", "PRESTASI",

    # Organization/Organizational Experience Variants
    "organization", "Organization", "ORGANIZATION",
    "organizational experience", "Organizational Experience", "ORGANIZATIONAL EXPERIENCE",
    "organisasi", "Organisasi", "ORGANISASI",
    "pengalaman organisasi", "Pengalaman Organisasi", "PENGALAMAN ORGANISASI",
    "ORGANIZATION EXPERIENCE", "organization experience", "ORGANIZATION EXPERIENCE",

    # Language Variants
    "languages", "Languages", "LANGUAGES",
    "bahasa", "Bahasa", "BAHASA",
    "kemampuan bahasa", "Kemampuan Bahasa", "KEMAMPUAN BAHASA",

    # Others
    "additional information", "Additional Information", "ADDITIONAL INFORMATION",
    "informasi tambahan", "Informasi Tambahan", "INFORMASI TAMBAHAN",
    "summary", "Summary", "SUMMARY",
    "pengalaman volunteer", "Pengalaman Volunteer", "PENGALAMAN VOLUNTEER",
    "volunteer experience", "Volunteer Experience", "VOLUNTEER EXPERIENCE",
    "volunteering", "Volunteering", "VOLUNTEERING",
    "Skills, Achievements & Other Experience", "skills, achievements & other experience", "SKILLS, ACHIEVEMENTS & OTHER EXPERIENCE",
    "skills, achievements & other experience", "Skills, Achievements & Other Experience", 
    "WORK EXPERIENCE / PROJECT TASKS",
    "Organisational Experience", "Organisational experience", "ORGANISATIONAL EXPERIENCE",
    "RINGKASAN PROFESIONAL", "COURSE & CERTIFICATION", "PROFILE",
    "COURSE CERTIFICATES","Awards, Achievements & Other Experience",
    "Skills and Language","Relevant Courses & Achievements","INTERNSHIP & Working EXPERIENCE",
    "Skill & Other Experience","CERTIFICATE OF COMPLETION"," PROJECTS DURING STUDY","SERTIFIKAT & PELATIHAN","Sertifikat & Pelatihan","sertifikat & pelatihan",
    "RELATED EXPERIENCES"
]


def build_section_regex(headers):
    """Bangun regex dinamis buat deteksi semua header"""
    joined = "|".join(re.escape(h) for h in headers)
    return re.compile(rf"(?i)(?:^|\n)\s*({joined})\s*(?:\n|:)", re.IGNORECASE)
def is_ats_friendly(text: str) -> bool:
    """
    Cek apakah CV kemungkinan ATS-friendly berdasarkan beberapa kriteria sederhana:
    - Punya setidaknya 3 section yang dikenali
    - Tidak terlalu banyak karakter non-alfanumerik (misal: >30% simbol)
    """

    sections = extract_all_sections(text, SECTION_HEADERS)
    if len(sections) < 3:
        return False

    non_alnum = sum(1 for c in text if not c.isalnum() and not c.isspace())
    if non_alnum / len(text) > 0.3:
        return False

    return True


def extract_all_sections(text, headers):
    """
    Ekstrak semua section dari teks CV berdasarkan daftar header.
    Hasilnya: dict {header: isi section}
    """
    pattern = build_section_regex(headers)
    matches = list(pattern.finditer(text))
    sections = {}

    for i, match in enumerate(matches):
        header = match.group(1).strip().lower()
        start = match.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        block = text[start:end].strip()
        sections[header] = block

    return sections

def extract_birthday(text: str) -> str:
    """
    Deteksi tanggal lahir dari CV.
    Format umum yang didukung:
    - 01-01-2000 / 01/01/2000
    - 1 Januari 2000
    - Tanggal Lahir: 01-01-2000
    """
    lines = [l.strip() for l in text.splitlines() if l.strip()]

    # Pola regex untuk berbagai format tanggal
    date_patterns = [
        r"\b\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b",  # 01-01-2000 atau 1/1/00
        r"\b\d{1,2}\s+(Januari|Februari|Maret|April|Mei|Juni|Juli|Agustus|September|Oktober|November|Desember)\s+\d{4}\b",  # 1 Januari 2000
        r"Tanggal\s+Lahir[:\-]?\s*\d{1,2}[-/]\d{1,2}[-/]\d{2,4}",  # Tanggal Lahir: 01-01-2000
        r"Tanggal\s+Lahir[:\-]?\s*\d{1,2}\s+(Januari|Februari|Maret|April|Mei|Juni|Juli|Agustus|September|Oktober|November|Desember)\s+\d{4}"  # Tanggal Lahir: 1 Januari 2000
    ]

    for line in lines:
        # skip baris yang isinya email, phone, link
        if any([
            re.search(r"@[a-zA-Z0-9.-]+\.[a-z]{2,}", line),
            re.search(r"\+?\d[\d\s-]{7,}", line),
            "linkedin" in line.lower(),
            "github" in line.lower(),
            extract_all_sections(text, SECTION_HEADERS)

        ]):
            continue

        # cek semua pola tanggal
        for pattern in date_patterns:
            match = re.search(pattern, line, re.IGNORECASE)
            if match:
                return match.group(0).replace("Tanggal Lahir", "").strip(": ").strip()

    return "Unknown"

def extract_address_and_domicile(text: str) -> dict:
    """
    Ekstraksi alamat & domisili (hybrid horizontal & vertical) dengan filtering kuat.
    Optimasi untuk berbagai format ATS: dot separators, pipes, commas, bullets.
    """
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    
    # ⚙️ Pre-filter: hanya keep baris pendek dan bukan paragraf
    filtered_lines = [
        l for l in lines[:15]
        if len(l) < 90 and l.count(",") <= 3
        and not re.search(r"\b(menjadi|melakukan|memiliki|pengalaman|berpengalaman|berkontribusi|menguasai|bersemangat)\b", l.lower())
    ]

    address = None
    domicile = None

    # Kata kunci diperluas
    address_keywords = ["jalan", "jl", "gg", "gang", "no", "rt", "rw", "blok", "kelurahan", "desa", "perumahan", "perum"]
    domicile_keywords = ["kota", "kabupaten", "provinsi", "kec", "kecamatan", "indonesia"]

    # Gabungkan baris atas dengan normalisasi separator
    top_text = " | ".join(filtered_lines[:8])

    # 🔍 OPTIMASI: Deteksi domisili dengan multi-pattern matching
    def extract_domicile_aggressive(txt: str) -> str:
        patterns = [
            # Format: "Kota Jakarta" atau "Provinsi Jawa Barat"
            r"(Kota|Kabupaten|Provinsi|Kec(?:\.|amatan)?)\s+([A-Z][a-zA-Z\s]+?)(?=\s*[|•,.\n]|,?\s*Indonesia|$)",
            # Format: "Jakarta, Indonesia" atau "Bandung, Jawa Barat"
            r"\b([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)?)\s*,\s*(Indonesia|Jawa\s+\w+|Sumatera\s+\w+|Kalimantan\s+\w+|Sulawesi\s+\w+|Bali|Papua|Maluku)",
            # Format: nama kota standalone di akhir segment
            r"(?:^|\||•|,\s*)([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)?)\s*(?:,\s*)?Indonesia",
            # Format sederhana: kota major Indonesia
            r"\b(Jakarta|Bandung|Surabaya|Medan|Semarang|Makassar|Palembang|Tangerang|Depok|Bekasi|Bogor|Yogyakarta|Malang|Denpasar|Balikpapan|Pontianak|Manado|Padang|Pekanbaru|Banjarmasin|Samarinda|Palangkaraya|Ambon|Jayapura|Kupang|Mataram|Serang|Cirebon|Tasikmalaya|Sukabumi|Cimahi|Solo|Surakarta|Magelang|Salatiga|Purwokerto|Tegal|Cilacap|Banda Aceh|Lhokseumawe|Binjai|Sibolga|Pematangsiantar|Bukittinggi|Jambi|Bengkulu|Palopo|Parepare|Kendari|Gorontalo|Manokwari|Ternate|Tidore|Sorong|Tarakan|Bontang|Probolinggo|Pasuruan|Sidoarjo|Mojokerto|Blitar|Kediri|Semarang Barat|Semarang Timur|Pontianak Kota|Denpasar Barat|Denpasar Timur)(?:\s+\w+)?(?:\s*,\s*Indonesia)?",
        ]
        
        for pattern in patterns:
            match = re.search(pattern, txt, re.IGNORECASE)
            if match:
                result = match.group(0).strip()
                # Clean up trailing punctuation
                result = re.sub(r'[|•]+$', '', result).strip()
                if len(result) > 3:
                    return result
        return None

    # Split dengan berbagai separator (enhanced)
    # Support: pipe |, bullet •, dot separator ". ", comma in context
    separator_pattern = r'[|\n•]|(?<![a-z])\.\s+(?=[A-Z])|(?<=\d)\.\s+(?=[A-Z])'
    parts = re.split(separator_pattern, top_text)

    for part in parts:
        p = part.strip()
        low = p.lower()

        if any((
            "http" in low or "www." in low,
            "linkedin" in low or "github" in low,
            re.search(r"@[a-zA-Z0-9._%+-]+\.[a-z]{2,}", p),
            re.search(r"\+?\d[\d\s\-().]{7,}", p),  # phone numbers
            len(p) < 3 or len(p) > 140,
            # 🚫 Block judul paper/artikel/project
            any(noise in low for noise in ["using", "analysis", "implementation", "development", "application", "system", "method", "approach", "study", "research", "performance", "optimization", "design", "model"]),
            "(" in p and ")" in p,  # Format paper/citation
        )):
            continue

        # 🎯 DOMICILE DETECTION (prioritas tinggi)
        if not domicile:
            # Cek apakah part ini mengandung domicile
            temp_domicile = extract_domicile_aggressive(p)
            if temp_domicile:
                domicile = temp_domicile
                continue
            
            # Fallback: keyword-based
            if any(k in low for k in domicile_keywords):
                # Validasi: bukan bagian dari alamat lengkap
                if not any(k in low for k in ["jalan", "jl.", "gang", "rt", "rw"]):
                    domicile = p
                    continue

        # ADDRESS DETECTION
        if not address and any(k in low for k in address_keywords):
            address = p
            continue

    # 🔄 FALLBACK: Global search di seluruh text
    if not domicile:
        domicile = extract_domicile_aggressive(text)
    
    if not domicile:
        # Last resort: cari pattern geografis
        m = re.search(
            r"(?:Kota|Kabupaten|Provinsi|Kec)\s+[A-Z][a-zA-Z\s]+?(?:,\s*Indonesia)?",
            text,
            re.IGNORECASE
        )
        if m:
            domicile = m.group(0).strip()

    if not address:
        m2 = re.search(
            r"\b(Jl\.?|Jalan|Gang|Gg\.?|No\.?|RT|RW)\b[^\n,]{3,100}",
            text,
            re.IGNORECASE
        )
        if m2:
            address = m2.group(0).strip()

    # ✅ VALIDATION dengan cleaning
    def clean_and_validate_domicile(s: str) -> str:
        if not s:
            return None
        # Remove trailing dots, pipes, bullets
        s = re.sub(r'[|•.]+$', '', s).strip()
        # Remove leading dots/pipes
        s = re.sub(r'^[|•.]+', '', s).strip()
        # Validasi: harus ada keyword atau "Indonesia"
        if any(k in s.lower() for k in domicile_keywords) or "indonesia" in s.lower():
            return s
        return None

    def clean_and_validate_address(s: str) -> str:
        if not s:
            return None
        s = re.sub(r'[|•]+$', '', s).strip()
        if any(k in s.lower() for k in address_keywords):
            return s
        return None

    domicile = clean_and_validate_domicile(domicile)
    address = clean_and_validate_address(address)

    return {
        "address": address or "Unknown",
        "domicile": domicile or "Unknown"
    }
def extract_resume_summary(text: str) -> str:
    """
    Ambil bagian resume/summary dari CV walaupun tanpa header.
    Logika baru (v4):
    - Fokus di 10–25 baris pertama (karena biasanya summary di atas)
    - Lewati baris berisi email, nomor HP, URL, LinkedIn, atau kota
    - Deteksi paragraf naratif (kalimat panjang, bukan bullet)
    - Stop kalau mulai muncul kata seperti EXPERIENCE, EDUCATION, SKILLS
    """

    lines = [l.strip() for l in text.splitlines() if l.strip()]
    top_part = lines[:125]

    section_keywords = [
        "experience", "education", "skills", "projects", "certification",
        "activities", "volunteer", "organization", "training", "PROFILE","Tentang Saya"
    ]

    summary_lines = []
    paragraph_started = False

    for line in top_part:
        lower = line.lower()

        # Skip baris yang mengandung kontak atau link
        if any([
            re.search(r"\b\d{5,}\b", line),  # zip code or long number
            re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-z]{2,}", line),
            re.search(r"(\+?\d[\d\s-]{8,15})", line),
            "linkedin.com" in lower or "github.com" in lower,
            re.match(r"^[A-Z\s,]+$", line)  # all caps or location line
        ]):
            continue

        # Kalau line mengandung section berikutnya, stop
        if any(word in lower for word in section_keywords):
            break

        # Skip bullet
        if line.startswith(("•", "-", "*", "—")):
            continue

        # Kalau ini naratif (panjang dan berisi kata kerja/subjektif)
        if len(line.split()) > 6:
            paragraph_started = True
            summary_lines.append(line)
        elif paragraph_started:
            # kalau udah mulai tapi line kosong, stop
            break

    summary = " ".join(summary_lines).strip()
    return summary if len(summary.split()) > 10 else "No summary found"

# === Bagian Metadata Dasar (nama, email, dll) ===

def parse_info(text: str, pdf_path: str = None) -> dict:
    # --- Ambil Info Umum ---
    email = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-z]{2,}", text)
    phone = re.search(r"(\+?\d[\d\s-]{8,15})", text)
    linkedin = re.search(r"(https?://(www\.)?linkedin\.com/in/[a-zA-Z0-9_-]+)", text)

    # storage = load_storage()
    # keywords = storage.get("skills_keywords", [])
    # skills = []
    # for word in keywords:
    #     if re.search(rf"\b{re.escape(word)}\b", text, re.IGNORECASE):
    #         skills.append(word)

    # --- Deteksi Nama Berdasarkan Font Terbesar (kalau PDF) ---
    name = "Unknown"
    if pdf_path:
        with pdfplumber.open(pdf_path) as pdf:
            first_page = pdf.pages[0]
            chars = first_page.chars
            if chars:
                lines = {}
                for c in chars:
                    y = round(c["top"], 1)
                    lines.setdefault(y, []).append(c)

                merged = []
                for y, chars in lines.items():
                    text_line = "".join(c["text"] for c in sorted(chars, key=lambda c: c["x0"]))
                    font_size = max(c["size"] for c in chars)
                    merged.append({"text": text_line.strip(), "size": font_size})

                biggest = max(merged, key=lambda x: x["size"], default={"text": "Unknown"})
                name = biggest["text"]


    # --- Ambil Semua Section Secara Dinamis ---
    sections = extract_all_sections(text, SECTION_HEADERS)
    domicile = extract_address_and_domicile(text)
    summary = extract_resume_summary(text)
    birthday = extract_birthday(text)
    ats_friendly =is_ats_friendly(text)

    return {
        "name": name,
        "email": email.group(0) if email else None,
        "phone": phone.group(0) if phone else None,
        "linkedin": linkedin.group(0) if linkedin else None,
        "domisili": domicile,
        "summary": summary,
        # "skill minimal":skill,
        "birthday": birthday,
        "sections": sections,
        "ats_friendly": ats_friendly,
    }


def extract_docx_with_font(path):
    doc = Document(path)
    extracted = []

    for para in doc.paragraphs:
        for run in para.runs:
            text = run.text.strip()
            if not text:
                continue

            # Ambil font size (dalam EMU → bisa None kalau pakai default style)
            font_size = run.font.size.pt if run.font.size else None
            bold = run.bold
            italic = run.italic

            extracted.append({
                "text": text,
                "size": font_size,
                "bold": bold,
                "italic": italic
            })

    return extracted


def secure_name(name: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]", "_", name)
